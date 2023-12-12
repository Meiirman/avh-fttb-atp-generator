from datetime import datetime
import json
import os
import re
import traceback

import pandas as pd
import tkinter as tk
from tkinter import messagebox, filedialog
from bs4 import BeautifulSoup
import openpyxl
from num2words import num2words
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from docx.enum.section  import WD_ORIENT # type: ignore

from scripts.models import Project


def log_text(text) -> None:
    global was_error
    was_error = True
    with open('media/logs.txt', 'a', encoding="utf-8") as log_file:
        current_datetime = datetime.now()
        formatted_datetime = current_datetime.strftime('%Y-%m-%d %H:%M:%S')
        log_file.write(f'{formatted_datetime}: {text}\n')
        
def get_work_folder() -> str | None:
    folder_path: str
    with open("config/work_folder.json", "r") as f:
        folder_path = json.load(f)["folder_path"]
    if folder_path != "":
        return folder_path
    else:
        send_message("Укажите путь к рабочей папке")
        return None


def get_smeta_data(smeta_path) -> dict:
    # Чтение Excel-файла
    df = pd.read_excel(smeta_path, header=None)
    
    TABLES: list[dict[dict]] = []
    table_started = False
    table_index = -1

    for index, value in enumerate(df[0]):

        print(value)
        
        if not isinstance(value, int) and table_started:
            table_started = False
            continue 
        
        elif not isinstance(value, int):
            continue

        elif value == 1:
            name = ""
            if "налог" in df[0][index-1].lower():
                name = ""
            elif "№" in df[0][index-1]:
                name = df[0][index-2]
            else:
                
                name = df[0][index-1]

            TABLES.append({
                "name" : name,
                "table" : []
            })
            table_index += 1
            table_started = True
            TABLES[table_index]["table"].append(
                {
                    "N" : df[0][index],
                    "D" : df[2][index],
                    "M" : df[3][index],
                    "C" : df[4][index],
                    "T" : "-",
                }
            )

        elif isinstance(value, int) and table_started:
            TABLES[table_index]["table"].append({
                    "N" : df[0][index],
                    "D" : df[2][index],
                    "M" : df[3][index],
                    "C" : df[4][index],
                    "T" : "-",
                })
    # print(json.dumps(TABLES, ensure_ascii=False, indent=4))
    # quit()
    
    # # Находим начало таблицы с помощью поиска ключевого слова
    # start_index = 0
    # try: start_index = df.index[df.apply(lambda row: 'Перенос шкафа Энергомера' in ' '.join(map(str, row)), axis=1)].tolist()[0]
    # except: start_index = df.index[df.apply(lambda row: 'СМЕТА' in ' '.join(map(str, row)), axis=1)].tolist()[0] + 1

    # # Находим конец таблицы с помощью поиска ключевого слова
    
    # try: end_index = df.index[df.apply(lambda row: 'Итого работа и материалы'.lower() in ' '.join(map(str, row)).lower(), axis=1)].tolist()[0]
    # except: end_index = df.index[df.apply(lambda row: 'Итого материалы и работа'.lower() in ' '.join(map(str, row)).lower(), axis=1)].tolist()[0]

    # # Выбираем подтаблицу с интересующими данными
    # table_data = df.iloc[start_index:end_index+1, :]

    # # Преобразовываем в двумерный массив
    # result_data = table_data.values.tolist()

    # result_data_items = []
    # result_data_works = []
    # is_work = False
    # indexes = {
    #     "№ п/п" : 0,
    #     "наименование" : 0,
    #     "ед. изм." : 0,
    #     "кол-во" : 0
    # }
    # result_row = result_data[2]
    # for i in result_data:
    #     if "ед. изм." in " ".join(list(map(str, i[:2]))):
    #         result_row = i
    #         break

    # for index, cell in enumerate(result_row):
    #     if f'{cell}'.lower() in "№ п/п".lower():
    #         indexes["№ п/п"] = index
    #     if f'{cell}'.lower() in "наименование".lower():
    #         indexes["наименование"] = index
    #     if f'{cell}'.lower() in "ед. изм." .lower():
    #         indexes["ед. изм." ] = index
    #     if f'{cell}'.lower() in "кол-во".lower():
    #         indexes["кол-во"] = index


    # for row in result_data[3:]:
    #     if "работа".lower() in " ".join(list(map(str, row[:2]))).lower():
    #         is_work = True
    #         continue

    #     if "Итого".lower() not in " ".join(list(map(str, row[:2]))).lower():
    #         if is_work:
    #             result_data_works.append({
    #                 "N" : row[indexes["№ п/п"]],
    #                 "D" : row[indexes["наименование"]],
    #                 "M" : row[indexes["ед. изм."]],
    #                 "C" : row[indexes["кол-во"]],
    #                 "T" : "-",
    #             })
    #         else:
    #             result_data_items.append({
    #                 "N" : row[indexes["№ п/п"]],
    #                 "D" : row[indexes["наименование"]],
    #                 "M" : row[indexes["ед. изм."]],
    #                 "C" : row[indexes["кол-во"]],
    #                 "T" : "-",
    #             })

    # result_dict = {
    #     "SMETA_ITEMS_TABLE" : result_data_items,
    #     "SMETA_WORKS_TABLE" : result_data_works
    # }

    return TABLES



def set_work_folder(folder_path):
    with open("config/work_folder.json", "w") as f:
        json.dump({"folder_path": folder_path}, f, ensure_ascii=False)
    send_message("Новое местоположение рабочей папки: \"" + folder_path + "\"")


def send_message(message):
    root = tk.Tk()
    root.withdraw()
    messagebox.showinfo("Формировать отчет", message)


def browse_folder(entry_var : tk.StringVar) -> None:
    folder_selected = filedialog.askdirectory()
    entry_var.set(folder_selected)
    set_work_folder(folder_selected)



def get_rvr_orders(project: Project) -> dict:
    folder = get_work_folder()
    if folder == None:
        return {"status" : -1}
    
    # получить все файлы в папке
    files = os.listdir(folder)
    excel_file_path = None
    if files:
        for file in files:
            if file.endswith(('.xlsx')):
                excel_file_path = folder + "/" + file

        workbook = openpyxl.load_workbook(excel_file_path)

        # Получаем активный лист (первый лист)
        sheet = workbook.active

        # Получаем значение из ячейки A1            

        # Получаем значение из ячейки A3
        value_A3 = sheet['A3'].value

        if value_A3:
            value_A3 = value_A3.split("в ")
        

        # Выводим значение
        print(f"Значение в ячейке A3: {value_A3}")

        # Закрываем файл Excel
        workbook.close()
    

    else:
        send_message("В рабочей папке нет файлов")
        return {"status" : -1}
    pass





def get_orders(project: Project) -> dict:
    # открыть папку
    folder = get_work_folder()
    if folder == None:
        return {"status" : -1}
    
    # получить все файлы в папке
    files = os.listdir(folder)

    # найти html файлы
    html_file_path = None
    if files:
        for file in files:
            if file.endswith(('.html')):
                html_file_path = folder + "/" + file
    else:
        send_message("В рабочей папке нет файлов")
        return {"status" : -1}

    if html_file_path == None:
        send_message("В папке нет html файла")
    
    # преобразовать html в правильный формат
    html_file = open(html_file_path, 'r', encoding='utf-8') # type: ignore
    print(html_file_path)
    html_content = replace_p_tags_with_br(html_file.read())

    # преобразовать html 
    soup = BeautifulSoup(html_content, 'html.parser')
    body = soup.find('body')
    result = []


    try:
        dogovor_data = body.findChildren(recursive=False)[0].find_all("table")[0].find("get_dogovor_data").text # type: ignore
        main_tables = body.findChildren(recursive=False)[0].find_all("table")[1:] # type: ignore
        vedomost_texts = [ row for row in body.findChildren(recursive=False)[0].find("td").get_text().split('\n') if "ВЕДОМОСТЬ исполнения работ" in row] # type: ignore
        print(vedomost_texts)
        
        multi_BS_NUMBER = []
        multi_BS_NAME = []
        multi_BS_ADDRESS = []
        multi_ORDER_REGION = []
        multi_ORDER_MANAGER = []
        multi_TOTAL_SUMM = []
        multi_TOTAL_NDS = []
        multi_TOTAL_SUMM_NDS = []
        multi_TOTAL_SUMM_NDS_WORD = []
        multi_ORDER_NUMBER = []
        multi_ORDER_DATE = []
        multi_TABLE = []

        const_ORDER_DOGOVOR_NUMBER = get_ORDER_DOGOVOR_NUMBER(dogovor_data) # DONE
        const_ORDER_DOGOVOR_DATE = get_ORDER_DOGOVOR_DATE(dogovor_data) # DONE

        for text in vedomost_texts:
            multi_BS_NUMBER.append(get_BS_NUMBER(text, f'{html_file}')) # DONE
            multi_BS_NAME.append(get_BS_NAME(text)) # DONE
            multi_BS_ADDRESS.append(get_BS_ADDRESS(text)) # DONE

        
        multi_ORDER_REGION = [regions[iii]["reg_name"] for iii in [i.text.strip() for i in soup.find_all("region_code")]] # DONE
        multi_ORDER_MANAGER = [regions[iii]["reg_resp_name"] for iii in [i.text.strip() for i in soup.find_all("region_code")]] # DONE
        multi_ORDER_MANAGER_POSITION = [regions[iii]["reg_resp_position"] for iii in [i.text.strip() for i in soup.find_all("region_code")]] # DONE

        for i in range(1, len(main_tables), 2):
            multi_TOTAL_SUMM.append(get_TOTAL_SUMM(main_tables[i])) # DONE
            multi_TOTAL_NDS.append(get_TOTAL_NDS(main_tables[i])) # DONE
            multi_TOTAL_SUMM_NDS.append(get_TOTAL_SUMM_NDS(main_tables[i])) # DONE
        
        for i in multi_TOTAL_SUMM_NDS:
            multi_TOTAL_SUMM_NDS_WORD.append(get_TOTAL_SUMM_NDS_WORD(i, num2words(int(i.strip().replace(" ", "").replace(",", ".").split(".")[0]), lang='ru'), f'{i.strip().replace(" ", "").replace(",", ".").split(".")[1]}'))
            multi_ORDER_NUMBER.append("")
            multi_ORDER_DATE.append("")
        
        for i in range(0, len(main_tables), 2):
            multi_TABLE.append(get_TABLE(main_tables[i])) # DONE

        
        for i, e in enumerate(multi_TABLE):
            data = {
                "BS_NUMBER" : multi_BS_NUMBER[i], 
                "BS_NAME" : multi_BS_NAME[i], 
                "BS_ADDRESS" : multi_BS_ADDRESS[i], 
                "ORDER_REGION" : multi_ORDER_REGION[i], 
                "ORDER_MANAGER" : multi_ORDER_MANAGER[i], 
                "ORDER_NUMBER" : multi_ORDER_NUMBER[i], 
                "ORDER_DATE" : multi_ORDER_DATE[i], 
                "TOTAL_SUMM" : multi_TOTAL_SUMM[i], 
                "TOTAL_NDS" : multi_TOTAL_NDS[i], 
                "TOTAL_SUMM_NDS" : multi_TOTAL_SUMM_NDS[i], 
                "TOTAL_SUMM_NDS_WORD" : multi_TOTAL_SUMM_NDS_WORD[i], 
                "ORDER_DOGOVOR_NUMBER" : const_ORDER_DOGOVOR_NUMBER, 
                "ORDER_DOGOVOR_DATE" : const_ORDER_DOGOVOR_DATE, 
                "TABLE" : multi_TABLE[i], 
                "ORDER_MANAGER_POSITION" : multi_ORDER_MANAGER_POSITION[i], 
                "TYPE_OF_WORK" : get_TYPE_OF_WORK(f'{html_file}')
            }
            result.append(data)
    except:
        log_text(traceback.format_exc())
        if project.show_errors_window:
            send_message("Произошла ошибка\n" + traceback.format_exc())
    return {"result" : result, "status" : 0}


def replace_p_tags_with_br(html_content):
    html_content = html_content.replace("style=\"font-size:0.12in;\"", "")
    html_content = html_content.replace("align=\"left\"", "")
    html_content = html_content.replace("align=\"center\"", "")
    html_content = html_content.replace("<br >", "")
    html_content = html_content.replace("<br>", "")
    html_content = html_content.replace("<p", "<br")
    html_content = html_content.replace("<b", "<br")
    html_content = html_content.replace("</p>", "")
    html_content = html_content.replace("brr", "br")
    html_content = html_content.replace("</br>", "")
    html_content = html_content.replace("</b>", "")
    html_content = html_content.replace("</b>", "")
    html_content = html_content.replace("<center>", "")
    html_content = html_content.replace("</center>", "")
    html_content = html_content.replace("<br >", "")
    html_content = html_content.replace("<br>", "")
    html_content = html_content.replace("<br/>", "")
    html_content = html_content.replace("brody", "body")
    html_content = html_content.replace("\n", "")
    html_content = html_content.replace("Итого стоимость работ", "\nИтого стоимость работ")
    html_content = html_content.replace("Всего общая стоимость работ", "\nВсего общая стоимость работ")
    html_content = html_content.replace("НДС 12%: ", "\nНДС 12%: ")
    html_content = html_content.replace("Номер заказа:", "\nНомер заказа:")
    html_content = html_content.replace("<", "\n<")
    html_content = html_content.replace("Регион: [", "\nРегион: [<region_code>")
    html_content = html_content.replace("] Номер Заявки", "</region_code>]\n Номер Заявки")
    
    html_content_x = html_content.split("\n")
    for i, e in enumerate(html_content_x):
        if "Итого стоимость работ" in e:
            html_content_x[i] = "<itogo_word>" + e + "</itogo_word>" 
        if "Всего общая стоимость работ" in e:
            html_content_x[i] = "<itogo_total_word>" + e + "</itogo_total_word>" 
        if "НДС 12%: " in e:
            html_content_x[i] = "<NDC_word>" + e + "</NDC_word>" 
        if "к рамочному договору" in e:
            html_content_x[i] = "<get_dogovor_data>" + e + "</get_dogovor_data>" 

    html_content = "\n".join(html_content_x)
    return html_content

def get_types_of_works(): 
    with open('config/типы работ.json', 'r', encoding="utf-8") as file:
        return json.load(file)

    
def get_TYPE_OF_WORK(file_name):
    return ""
    types_of_works = get_types_of_works()

    return types_of_works[re.findall(r'\((.*?)\)', file_name.split("\\")[-1])[0]]
    
    
def get_FILE_NAME(ATP_OR_AVR, BS_NAME, TYPE_OF_WORK):
    variants = {
        "АТП" : {
            "демонтажных работ" : f"АТП_ДМР_{BS_NAME}_",
            "монтажных работ" : f"АТП_МР_{BS_NAME}_",
            "строительных работ" : f"{BS_NAME}_АТП_",
            "электро-монтажных работ" : f"АТП_ЭМР_{BS_NAME}_",
        },
        "АВР" : {
            "демонтажных работ" : f"АВР_ДМР_{BS_NAME}_",
            "монтажных работ" : f"АВР_МР_{BS_NAME}_",
            "строительных работ" : f"АВР_СР_{BS_NAME}_",
            "электро-монтажных работ" : f"АВР_ЭМР_{BS_NAME}_",
        }
    }

    # return variants[ATP_OR_AVR][TYPE_OF_WORK]
    return "FTTB АТП " + BS_NAME


def get_regions(): 
    with open('config/regions.json', 'r', encoding="utf-8") as file:
        return json.load(file)


regions = get_regions()


def get_TABLE(table):
    TABLE = []
    if table:
        rows = table.find_all('tr')
        for row in rows:
            cells = row.find_all('td')
            row_list = []
            for cell in cells:
                row_list.append(cell.text.strip())
            try:
                int(row_list[0])
                i = row_list
                TABLE.append({"N" : i[0], "P" : i[1], "D" : i[2], "M" : i[3], "C" : i[4], "T" : i[6], "S" : i[5]})
            except:
                # traceback.print_exc()
                pass        
    # print(TABLE)
    return TABLE

    
def get_BS_NUMBER(text, file_name):
    try: return [i for i in file_name.split("_") if "БС№" in i ][0]
    except:
        try: return [i for i in text.split() if "БС№" in i][0] 
        except: return ""
            

    
def get_BS_NAME(text):
    print(text)
    name = text
    try: name = text.split("ВЕДОМОСТЬ исполнения работ ")[1]
    except: pass

    try: name = name.split(" в ")[0]
    except: pass

    # words = text.split("\"")
    return name 

    
def get_BS_ADDRESS(text):
    return "г." + f'{text.split("ВЕДОМОСТЬ исполнения работ")[1].split("г.")[1]}'

    
def get_ORDER_REGION(soup):
    try:
        comment_elements = soup.find_all(string=lambda text: "Vedomost" in text)
        comments_text = [comment.strip() for comment in comment_elements][0]
        extract_metadata(comments_text)['region'] # type: ignore
        return regions[extract_metadata(comments_text)['region']]["reg_name"] # type: ignore
    except:
        return "reg_name - Не известено"

    
def get_ORDER_MANAGER(soup):
    try:
        comment_elements = soup.find_all(string=lambda text: "Vedomost" in text)
        comments_text = [comment.strip() for comment in comment_elements][0]
        extract_metadata(comments_text)['region'] # type: ignore
        return regions[extract_metadata(comments_text)['region']]["reg_resp_name"] # type: ignore
    except:
        return "reg_resp_name - Не известено"


def get_ORDER_MANAGER_POSITION(soup):
    try:
        comment_elements = soup.find_all(string=lambda text: "Vedomost" in text)
        comments_text = [comment.strip() for comment in comment_elements][0]
        extract_metadata(comments_text)['region'] # type: ignore
        return regions[extract_metadata(comments_text)['region']]["reg_resp_position"] # type: ignore
    except:
        return "reg_resp_position - Не известено"

    

def get_ORDER_NUMBER(soup):
    ORDER_NUMBER = ""
    return ORDER_NUMBER

    
def get_ORDER_DATE(soup):
    ORDER_DATE = ""
    return ORDER_DATE

    
def get_TOTAL_SUMM(table):
    rows = table.find_all('tr')
    for row in rows:
        cells = row.find_all('td')
        if "Итого:" in cells[0].text.strip():
            return cells[1].text.strip()
    TOTAL_SUMM = ""
    return TOTAL_SUMM

    
def get_TOTAL_NDS(table):
    rows = table.find_all('tr')
    for row in rows:
        cells = row.find_all('td')
        if "НДС 12%:" in cells[0].text.strip():
            return cells[1].text.strip()
    TOTAL_NDS = ""
    return TOTAL_NDS

    
def get_TOTAL_SUMM_NDS(table):
    rows = table.find_all('tr')
    for row in rows:
        cells = row.find_all('td')
        if "учетом НДС:" in cells[0].text.strip():
            return cells[1].text.strip()
    TOTAL_SUMM_NDS = ""
    return TOTAL_SUMM_NDS

    
def get_TOTAL_SUMM_NDS_WORD(num, num_word, kopeiki):

    TOTAL_SUMM_NDS_WORD = f"Всего общая стоимость работ: {num} ( {num_word} ) тенге и {kopeiki} тиын"

    return TOTAL_SUMM_NDS_WORD

    
def get_ORDER_DOGOVOR_NUMBER(dogovor_data):
    xx = dogovor_data.split()
    for i in xx:
        if "№" in i and "№" != i:
            return i[1:]
    ORDER_DOGOVOR_NUMBER = ""
    return ORDER_DOGOVOR_NUMBER

    
def get_ORDER_DOGOVOR_DATE(dogovor_data):
    try:
        for i in dogovor_data.split():
            if "г" == i[len(i)-1]:
                return i[:-1]
    except:
        # traceback.print_exc()
        pass
    ORDER_DOGOVOR_DATE = ""
    return ORDER_DOGOVOR_DATE


def get__there_should_be_an_smeta_if_there_is_this_text():
    there_should_be_an_smeta_if_there_is_this_text: str
    with open("config/config.json", "r") as f:
        there_should_be_an_smeta_if_there_is_this_text = json.load(f)["there_should_be_an_smeta_if_there_is_this_text"]
    if there_should_be_an_smeta_if_there_is_this_text != "":
        return there_should_be_an_smeta_if_there_is_this_text
    else:

        # send_message("Укажите путь к рабочей папке")
        return ""


def get_have_smeta(order) -> bool:
    order_str = f"{order}"
    there_should_be_an_smeta_if_there_is_this_text = get__there_should_be_an_smeta_if_there_is_this_text()
    if there_should_be_an_smeta_if_there_is_this_text.lower() in order_str.lower():
        return True
    return False


def combine_docx(file1, file2, output_file, is_second=False, is_atp=False):
    doc1 = Document(file1)
    doc2 = Document(file2)

    docX = doc1

    docX.add_paragraph('')
    docX.add_paragraph('')
    # docX.add_page_break()



    for section in docX.sections:
        section.orientation = WD_ORIENT.PORTRAIT

    for element in doc2.element.body:
        docX.element.body.append(element)

    combined_doc = Document()
    combined_doc = docX

    for section in combined_doc.sections:
        section.page_width = Pt(510)
        section.page_height = Pt(728)

    # if len(combined_doc.tables) > 3:
    for index, table in enumerate(combined_doc.tables):
        table.autofit = True  # Отключаем автонастройку ширины столбцов

        if index > 3:
            col_lens = []
            for row in table.rows:
                col_lens.append(len(row.cells))

                num_cols = len(row.cells)
                for cell in row.cells:
                    cell.width = Pt(728/num_cols)

    combined_doc.save(output_file)  

    doc1.save(file1[:-5] + ".docx" )


def get_smeta(order):
     # открыть папку
    folder = get_work_folder()
    if folder == None:
        return {"status" : -1}
    
    # получить все файлы в папке
    files = os.listdir(folder)

    # найти xlsx файлы
    xlsx_files = []
    if files:
        for file in files:
            if file.endswith(('.xlsx')):
                xlsx_files.append(folder + "/" + file )  
        if len(xlsx_files) == 0:
            send_message("Для заказа требуется смета которую не нашел в папке. Пожалуйста добавьте смету в папку")
            return ""
    else:
        send_message("В рабочей папке нет файлов")
        return {"status" : -1}
    

    for file in xlsx_files:
        if order['BS_ADDRESS'][2:10] in file:
            return file

    send_message("В папке нет нужной сметы. Пожалуйста добавьте смету в папку")
    return ""


def ADD_END(typez, input_path, output_path, data):
    template = None
    if typez=="avr": template = DocxTemplate("templates/ШАБЛОН АВР END.docx")
    else: template = DocxTemplate("templates/ШАБЛОН АТП END.docx")

    # print(data)
    template.render(data)
    template.save("templates/ШАБЛОН WITH END.docx")

    combine_docx(input_path, "templates/ШАБЛОН WITH END.docx", output_path, True, typez == "АТП" )


def create_files(folder, data, tmpl_type, have_smeta=False, selected_date=None):
    data['R_T'] = selected_date.strftime("%d.%m.%Y")
    if " - " in data['BS_NAME']:
        BS_ADDRESSx=data['BS_ADDRESS']
        BS_ADDRESS = data['BS_ADDRESS'].split(" - ")
        BS_NAME = data['BS_NAME'].split(" - ")
        try: data['BS_ADDRESS'] = f'{BS_NAME[0]} - {BS_ADDRESS[0]}\n{BS_NAME[1]} - {BS_ADDRESS[1]}'
        except: data['BS_ADDRESS'] = BS_ADDRESSx
        
    smeta_path = ""
    if have_smeta:
        smeta_path = get_smeta(data)

    if "atp" in tmpl_type:
        data["WORK_NAME"] = data['BS_NAME']
        data["have_smeta"] = have_smeta
        template_ATP = DocxTemplate("templates/FTTB ШАБЛОН АТП.docx")    
        
        if smeta_path != "":
            data["SMETA_TABLES"] = get_smeta_data(smeta_path)
            # print(smeta_data)
            # data["SMETA_TABLES"] = smeta_data["SMETA_ITEMS_TABLE"]
            # data["SMETA_WORKS_TABLE"] = smeta_data["SMETA_WORKS_TABLE"]
            # combine_docx(output_path, smeta_path, output_path, is_second=False, is_atp=False)
        # ADD_END("atp", output_path, output_path, data)

        template_ATP.render(data)
        file_name__ATP = get_FILE_NAME("АТП", data['BS_NAME'], data['TYPE_OF_WORK'])
        output_path = folder + "/" + file_name__ATP + ".docx"
        template_ATP.save(output_path)
    

    if "avr" in tmpl_type:
        template_AVR = DocxTemplate("templates/ШАБЛОН АВР.docx")    
        template_AVR.render(data)
        file_name__AVR = get_FILE_NAME("АВР", data['BS_NAME'], data['TYPE_OF_WORK'])
        output_path = folder + "/" + file_name__AVR + ".docx"
        template_AVR.save(output_path)

        if smeta_path != "":
            combine_docx(output_path, smeta_path, output_path, is_second=False, is_atp=False)
        ADD_END("avr", output_path, output_path, data)
    